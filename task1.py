from math import log                #ver 0.1.3
from itertools import chain
from turtle import width

from docx.shared import Cm, Pt
from tabulate import tabulate

from utils.task1_utils import *
from utils.docx_utils import *




def task1_1(pn, listedLetters, document):
    document_output = []
    document_output.append("""1.1 Скласти шестизначне число, яке складається з отриманих за допомогою кодової таблиці кодів 1-ої, 2-ої та 8-ої літер прізвища. При цьому перші 3 цифри відповідають цілій частині числа, а останні - дробовій.
    Вважаючи це число десятковим, перевести його до шістнадцяткової, з шістнадцяткової до двійкової, з двійкової до вісімкової систем числення з точністю відповідно 3, 5 та 3 розрядів після коми.""")
    print("\n\n1.1")
    document_output.append(f"{listedLetters[0]} - {pn[0][0]}{pn[0][1]}, {listedLetters[1]} - {pn[1][0]}{pn[1][1]}, {listedLetters[7]} - {pn[7][0]}{pn[7][1]}")
    splitter()
    number = float(f"{pn[0][0]}{pn[0][1]}{pn[1][0]}.{pn[1][1]}{pn[7][0]}{pn[7][1]}") # for output
    document_output.append(f"Число: {number}₁₀")
    numberBD = int(f"{pn[0][0]}{pn[0][1]}{pn[1][0]}") # before dot
    numberAD = int(f"{pn[1][1]}{pn[7][0]}{pn[7][1]}") # after dot
    print(f"Десяткове число, трансформоване з персональних кодів - {number}")
    splitter()
    document_output.append("Переведення з десяткового числа до шіснадцяткового:")
    print("Переведення з десяткового числа до шіснадцяткового:")
    numberBD, co = decimal_to_hexadecimal_bd(numberBD)
    document_output += co
    print(f"Число до коми дорівнює {numberBD}")
    numberAD, co = decimal_to_hexadecimal_ad(numberAD)
    document_output += co
    print(f"Число після коми дорівнює {numberAD}")
    document_output.append(f"{number}₁₀ = {numberBD}.{numberAD}₁₆")
    number = str(f"{numberBD}.{numberAD}")
    prevnumber = number
    print(f"З'єднане число дорівнює {number}")
    splitter()
    document_output.append("Переведення з шіснадцяткового числа до двійкового:")
    print("Проведемо перетворення з шіснадцяткового числа до двійкового.")
    numberBD, numberAD, number, co = hexadecimal_to_binary(numberBD, numberAD)
    document_output += co
    document_output.append(f"{prevnumber}₁₆ = {number}₂")
    prevnumber = number
    print(f"Після очищення від непотрібних нулів та скорочення, число дорівнює {number}")
    splitter()
    document_output.append("Переведення з двійкового числа до вісімкового:")
    print("Проведемо перетворення з двійкового числа до вісімкового.")
    number, numberBD, numberAD, co = binary_to_octo(numberBD, numberAD)
    document_output += co
    document_output.append(f"{prevnumber}₂ = {number}₈")
    print(f"Після форматування, число дорівнює {number}")
    splitter()
    for i in range(len(document_output)):
        document.add_paragraph(document_output[i])
    document.add_page_break()
    

def task1_2(pn, listedLetters, document):
    document_output = []
    document_output.append("""1.2 Скласти шестизначне число, яке складається з отриманих за допомогою кодової таблиці кодів 1-ої, 2-ої та 8-ої літер прізвища. При цьому перші 3 цифри відповідають цілій частині числа, а останні - дробовій.
    Вважаючи це число шістнадцятковим, перевести його до десяткової, з шістнадцяткової до двійкової, з двійкової до вісімкової систем числення з точністю відповідно 3, 5 та 3 розрядів після коми.""")
    print("\n\n1.2")
    splitter()
    document_output.append(f"{listedLetters[0]} - {pn[0][0]}{pn[0][1]}, {listedLetters[1]} - {pn[1][0]}{pn[1][1]}, {listedLetters[7]} - {pn[7][0]}{pn[7][1]}")
    number = f"{pn[0][0]}{pn[0][1]}{pn[1][0]}.{pn[1][1]}{pn[7][0]}{pn[7][1]}"
    numberBD = f"{pn[0][0]}{pn[0][1]}{pn[1][0]}"
    numberAD = f"{pn[1][1]}{pn[7][0]}{pn[7][1]}"
    document_output.append(f"Число: {number}₁₆")
    prevnumber = number
    print(f"Шіснадцяткове число, трансформоване з персональних кодів - {number}")
    splitter()
    document_output.append("Переведення з шіснадцяткового числа до десяткового:")
    print("Проведемо перетворення з шіснадцяткового числа до десяткового.")
    number, numberBD, numberAD, co = hexadecimal_to_decimal(numberBD, numberAD)
    document_output += co
    document_output.append(f"{prevnumber}₁₆ = {number}₁₀")
    print(f"Число до коми дорівнює {numberBD}")
    print(f"Після проведеного округлення, число після коми дорівнює {numberAD}")
    print(f"З'єднане число дорівнює {number}")
    splitter()

    number = str(f"{pn[0][0]}{pn[0][1]}{pn[1][0]}.{pn[1][1]}{pn[7][0]}{pn[7][1]}")
    numberBD = str(f"{pn[0][0]}{pn[0][1]}{pn[1][0]}")
    numberAD = str(f"{pn[1][1]}{pn[7][0]}{pn[7][1]}")
    prevnumber = number
    document_output.append("Переведення з шіснадцяткового числа до двійкового:")
    print("Проведемо перетворення з шіснадцяткового числа до двійкового.")
    numberBD, numberAD, number, co = hexadecimal_to_binary(numberBD, numberAD)
    document_output += co
    document_output.append(f"{prevnumber}₁₆ = {number}₂")
    print(f"Після очищення від непотрібних нулів та скорочення, число дорівнює {number}")
    splitter()

    document_output.append("Переведення з двійкового числа до вісімкового:")
    prevnumber = number
    print("Проведемо перетворення з двійкового числа до вісімкового.")
    number, numberBD, numberAD, co = binary_to_octo(numberBD, numberAD)
    document_output += co
    document_output.append(f"{prevnumber}₂ = {number}₈")
    print(f"Після форматування, число дорівнює {number}")
    splitter()
    for i in range(len(document_output)):
        document.add_paragraph(document_output[i])
    document.add_page_break()



def task1_3(pn, document):
    print("\n\n1.3")
    number = int(f"{pn[0][0]}{pn[0][1]}{pn[1][0]}{pn[1][1]}{pn[7][0]}{pn[7][1]}")
    p = [2,3,5,7,11,13,17]
    P = 510510
    numberMod = []
    B = [0,0,0,0,0,0,0]
    sum = 0
    document_output = []
    document_output.append("Число: " + str(number))
    for i in range(7):
        document_output.append("p" + chr(int("832" + str(i+1))) + " = " + str(p[i]))
        print(document_output[i+1])

    document_output.append("P = p₁ * p₂ * p₃ * p₄ * p₅ * p₆ * p₇ = " + str(P))
    print(document_output[8])

    for i in range(7):
        numberMod.append(number % p[i])
        document_output.append(str(number) + " mod " + str(p[i]) + " = " + str(numberMod[i]))
        print(document_output[9+i])

    
    j = 16
    for i in range(7):
        x = 1
        while True:
            B[i] = int(x * (P/p[i]))
            document_output.append("B" + chr(int("832" + str(i+1))) + " = " + str(x) + " * " + str(P) + " : " + str(p[i]) + " = " + str(B[i]) + " (" + str(int(B[i] % p[i])) + ")")
            print(document_output[j])
            j+=1
            x+=1
            if B[i] % p[i] == 1:
                break
    
    for i in range(7):
        sum += numberMod[i] * B[i]

    document_output.append("( " + str(numberMod[0]) + " * " + str(B[0]) + " + " + str(numberMod[1]) + " * " + str(B[1]) + " + " + str(numberMod[2]) + " * " + str(B[2]) + " + " + str(numberMod[3]) + " * " + str(B[3])
    + " + " + str(numberMod[4]) + " * " + str(B[4]) + " + " + str(numberMod[5]) + " * " + str(B[5]) + " + " + str(numberMod[6]) + " * " + str(B[6]) + ") = " + str(sum) + " mod " + str(P) + " = " + str(sum % P))
    print(document_output[len(document_output)-1])

    document.add_paragraph("1.3 Скласти шестизначне число, яке складається з отриманих за допомогою кодової таблиці кодів 1-ої, 2-ої та 8-ої літер прізвища. Вважаючи це число десятковим, перевести його до системи числення залишкових класів із мінімально необхідною кількістю основ 2, 3, 5, 7, 11, ... . Після цього зробити зворотне переведення отриманого результату до десяткової системи числення.\n")

    for i in range(len(document_output)):
        document.add_paragraph(document_output[i])

    document.add_page_break()





def task1_4(listedLetters, pn, document):
    print("\n\n1.4")
    document.add_paragraph("""1.4 Виконати ефективне кодування визначених літер прізвища, при умові, що отримане за допомогою кодової таблиці число - десяткове і говорить про те, скільки разів у "повідомленні" зустрічається дана літера (при цьому, "повідомлення" складається всього з 8 обраних літер). 
Визначити ефективність проведеного кодування та порівняти її з ентропією джерела повідомлення і ефективністю рівномірного кодування, тобто, з випадком, коли довжина коду для кожної літери одна й та сама. За допомогою отриманих кодів скласти повідомлення, яке складається з визначених літер у тій послідовності, в якій вони зустрічаються у прізвищі. Визначити довжину (в бітах) повідомлення при ефективному і рівномірному кодуванні.""")
    n = 8 #number of symbols
    document.add_paragraph("")
    for i in range(8):
        p[i].sym = listedLetters[i]
        p[i].k = int(f"{pn[i][0]}{pn[i][1]}")

    totalK = 0
    
    for i in range(8):
        totalK += p[i].k

    totalPro = 1

    #p[0].pro = 0.3
    #p[1].pro = 0.19
    #p[2].pro = 0.12
    #p[3].pro = 0.11
    #p[4].pro = 0.1
    #p[5].pro = 0.08
    #p[6].pro = 0.06
    #p[7].pro = 0.04

    for i in range(8):
        p[i].pro = round(p[i].k / totalK,2);
        totalPro -= p[i].pro
        if i == 7:
            p[i].pro = round(p[i].pro + totalPro,2)

    dtp = displayRaw(n, p)
    firstTable = document.add_table(rows=9, cols=3)
    firstTable.autofit = False
    firstTable.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(8):
        for j in range(3):
            block = firstTable.rows[i].cells[j]
            set_cell_border(
                block,
                top={"sz": 5, "val": "single", "color": "#000000"},
                bottom={"sz": 5, "val": "single", "color": "#000000"},
                start={"sz": 5, "val": "single", "color": "#000000"},
                end={"sz": 5, "val": "single", "color": "#000000"},
            )
            firstTable.rows[i].cells[j].text = str(dtp[j][i])

    set_cell_border(
        firstTable.rows[8].cells[1],
        top={"sz": 5, "val": "single", "color": "#000000"},
        bottom={"sz": 5, "val": "single", "color": "#000000"},
        start={"sz": 5, "val": "single", "color": "#000000"},
        end={"sz": 5, "val": "single", "color": "#000000"},
    )
    set_cell_border(
        firstTable.rows[8].cells[2],
        top={"sz": 5, "val": "single", "color": "#000000"},
        bottom={"sz": 5, "val": "single", "color": "#000000"},
        start={"sz": 5, "val": "single", "color": "#000000"},
        end={"sz": 5, "val": "single", "color": "#000000"},
    )
    firstTable.rows[8].cells[1].text = "Σ = " + str(totalK)
    firstTable.rows[8].cells[2].text = "Σ = 100%"

    for i in range(9):
        firstTable.rows[i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        firstTable.rows[i].height = Pt(14)
        for j in range(3):
            cell = firstTable.rows[i].cells[j]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if j == 0:
                cell.width = Cm(1)
            else:
                cell.width = Cm(2.2)

    print("\n\t\t", "Σ=" + str(totalK), "\t\t", 1,"\t",end='') #підсумовує деякі дані

    sortByProbability(n, p, ['000', '001', '010', '011', '100', '101', '110', '111'])
 
    for i in range(n):
        p[i].top = -1

    shannon(0, n - 1, p)
    dtp = displayCalculated(n, p)
    document.add_paragraph("")
    secondTable = document.add_table(rows=9, cols=6, style="Table Grid")
    secondTable.autofit = False
    secondTable.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i in range(9):
        for j in range(6):
            cell = secondTable.rows[i].cells[j]
            secondTable.rows[i].cells[j].text = str(dtp[j][i])
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if j == 0:
                cell.width = Cm(2)
            elif j == 1:
                cell.width = Cm(3)
            elif j == 2:
                cell.width = Cm(4)
            elif j == 3:
                cell.width = Cm(3)
            elif j == 4:
                cell.width = Cm(4.2)
            elif j == 5:
                cell.width = Cm(3)

    H = 0
    print("\n\nH = -(", end='')
    dtp = "H = -("
    for i in range(8):
        H += p[n-1-i].pro * log(p[n-1-i].pro,2)
        if i != 7:
            dtp += str(p[n-1-i].pro) + " * log2(" + str(p[n-1-i].pro) + ") + "
            print(str(p[n-1-i].pro) + " * log2(" + str(p[n-1-i].pro) + ") + ", end='')
        else:
            dtp += str(p[n-1-i].pro) + " * log2(" + str(p[n-1-i].pro) + "))=" + str(round(-H, 2))
            print(str(p[n-1-i].pro) + " * log2(" + str(p[n-1-i].pro) + "))=" + str(round(-H, 2)),end='')
    document.add_paragraph("")
    document.add_paragraph(dtp)
    H = -H
    lEffective = 0
    document.add_paragraph("L сер. не еф = 3")
    print("\nL сер. не еф = 3")
    dtp = "L сер. еф = "
    print("L сер. еф = ", end='')

    for i in range(8):
        k = 0
        for j in range(p[n-1-i].top+1):
            if(p[n-1-i].arr[j] == 1 or p[n-1-i].arr[j] == 0):
                k += 1
        lEffective += k * p[n-1-i].pro
        if i != 7:
            dtp += str(k) + " * " + str(p[n-1-i].pro) + " + "
            print(str(k) + " * " + str(p[n-1-i].pro) + " + ", end='')
        else:
            dtp += str(k) + " * " + str(p[n-1-i].pro) + " = " + str(round(lEffective, 2))
            print(str(k) + " * " + str(p[n-1-i].pro) + " = " + str(lEffective), end='')

    document.add_paragraph(dtp)
    
    if lEffective < H and H < 3:
        document.add_paragraph("L сер. еф < H < L сер. не еф")
        print("\n\nL сер. еф < H < L сер. не еф")
    elif lEffective > H and lEffective < 3:
        document.add_paragraph("H < L сер. еф < L сер. не еф")
        print("\n\nH < L сер. еф < L сер. не еф")
    elif lEffective > H and lEffective > 3:
        document.add_paragraph("H < L сер. не еф < L сер. еф")
        print("\n\nH < L сер. не еф < L сер. еф")

    effectiveOutput = ""
    notEffectiveOutput = ""

    for i in range(8):
        for j in range(8):
            if listedLetters[i] == p[j].sym:
                for x in range(p[j].top+1):
                    effectiveOutput += str(p[j].arr[x])
                notEffectiveOutput += str(p[j].notEffectiveCode)
        effectiveOutput += " "
        notEffectiveOutput += " "
    print("\n\n")
    document.add_paragraph(effectiveOutput + " - " + str(len(effectiveOutput) - effectiveOutput.count(' ')) + " біт ефективний код")
    print(effectiveOutput + "  " + str(len(effectiveOutput) - effectiveOutput.count(' ')) + "біт")
    document.add_paragraph(notEffectiveOutput + " - " + str(len(notEffectiveOutput) - notEffectiveOutput.count(' ')) + " біт не ефективний код")
    print(notEffectiveOutput + "  " + str(len(notEffectiveOutput) - notEffectiveOutput.count(' ')) + "біт")
    document.add_page_break()


def task1_5(ll, pn, document):
    document.add_paragraph("1.5 Для шістнадцяти розрядного двійкового коду (1ц1л)(2ц1л)(1ц8л)(2ц8л) сформувати код Геммінга (Hamming) і продемонструвати його реакцію на однократний збій. Результати подати у вигляді таблиці.")
    print("\n\n1.5")
    splitter()
    
    numberHex = f"{pn[0][0]}{pn[0][1]}{pn[7][0]}{pn[7][1]}"
    numberBin = bin(int(numberHex, 16))[2:]
    numberBin = full_form_bin(numberBin, 16)
    
    outStrYourNumbers = f"{ll[0]}: {numberHex[0:2]}; {ll[7]}: {numberHex[2:4]}; {numberHex}={numberBin}\n"
    document.add_paragraph(outStrYourNumbers)
    print(outStrYourNumbers)
    
    controlPositions = get_control_pos(numberBin)
    numberBinLst = convert_to_list(numberBin, controlPositions)
    numberLen = len(numberBinLst)
    
    # Binary indeces for check matrix
    binIndeces = []
    for i in range(1, numberLen+1):
        binIndeces.append(list(full_form_bin(bin(i)[2:], 5))) # remove constant "5"
    
    # Fill check matrix
    checkMatrix = []
    row = []
    for i in range(len(binIndeces[0])): # range(5)
        for j in range(len(binIndeces)):
            row.append(int(binIndeces[j][i]))
        checkMatrix.append(row.copy())
        row.clear()

    j = 0
    for i in range(1, numberLen+1):
        if i in controlPositions[j:]:
            row.append(f"k{i}")
            j += 1
        else:
            row.append(f"i{i}")
    
    checkMatrix.append(row.copy())
    row.clear()
    checkMatrix.append(numberBinLst.copy())
    
    # Calculate k1
    k1, outStrCalck1 = calc_k1(numberBinLst, 'k')
    # Calculate k2
    k2, outStrCalck2 = calc_k2(numberBinLst, 'k')
    # Calculate k4
    k4, outStrCalck4 = calc_k4(numberBinLst, 'k')
    # Calculate k8
    k8, outStrCalck8 = calc_k8(numberBinLst, 'k')
    # Calculate k16
    k16, outStrCalck16 = calc_k16(numberBinLst, 'k')

    checkMatrix[-1][0] = k1
    checkMatrix[-1][1] = k2
    checkMatrix[-1][3] = k4
    checkMatrix[-1][7] = k8
    checkMatrix[-1][15] = k16

    # Imitation of a failure bit
    posOfFailBit = random_fail_bit_pos(numberLen, controlPositions)
    numberBinLst[posOfFailBit-1] = int(not numberBinLst[posOfFailBit-1])

    # Calculate K1
    k1_, outStrCalcK1 = calc_k1(numberBinLst, 'K')
    # Calculate K2
    k2_, outStrCalcK2 = calc_k2(numberBinLst, 'K')
    # Calculate K4
    k4_, outStrCalcK4 = calc_k4(numberBinLst, 'K')
    # Calculate K8
    k8_, outStrCalcK8 = calc_k8(numberBinLst, 'K')
    # Calculate K16
    k16_, outStrCalcK16 = calc_k16(numberBinLst, 'K')
    
    # Find position of the failure bit
    posOfFailBit_ = f"{k16_^k16}{k8_^k8}{k4_^k4}{k2_^k2}{k1_^k1}"
    outStrCalcFailBitPos = "K ⊕ k = (K16 ⊕ k16)(K8 ⊕ k8)(K4 ⊕ k4)(K2 ⊕ k2)(K1 ⊕ k1) = " \
        f"({k16_} ⊕ {k16})({k8_} ⊕ {k8})({k4_} ⊕ {k4})({k2_} ⊕ {k2})({k1_} ⊕ {k1}) = {posOfFailBit_}"

    rows, cols = len(checkMatrix), len(checkMatrix[0])
    document.add_paragraph("Перевірочна матриця (таблиця)")
    table = document.add_table(rows=rows, cols=cols, style="Table Grid")
    table.autofit = False
    rowN = 0
    for row in checkMatrix:
        colN = 0
        for item in row:
            cell = table.rows[rowN].cells[colN]
            cell.text = str(item)
            cell.width = Cm(0.95)
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(10)
            colN += 1
        rowN += 1


    tableFailBit = document.add_table(rows=1, cols=cols)
    failBit = int(not checkMatrix[rows-1][posOfFailBit-1])
    tableFailBit.rows[0].cells[posOfFailBit-1].text = f"({failBit})"
    # Setting cell width and font size
    tableFailBit.autofit = False
    for row in tableFailBit.rows:
        for cell in row.cells:
            cell.width = Cm(0.95)
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(10)

    document.add_paragraph(f"Припустимо, що помилка сталася у біті i{posOfFailBit}.")
    document.add_paragraph()

    print_table(checkMatrix, posOfFailBit)

    print(checkMatrix) # test

    document.add_paragraph(outStrCalck1)
    print(outStrCalck1)
    document.add_paragraph(outStrCalck2)
    print(outStrCalck2)
    document.add_paragraph(outStrCalck4)
    print(outStrCalck4)
    document.add_paragraph(outStrCalck8)
    print(outStrCalck8)
    document.add_paragraph(outStrCalck16)
    print(outStrCalck16)
    document.add_paragraph()
    print("")
    document.add_paragraph(outStrCalcK1)
    print(outStrCalcK1)
    document.add_paragraph(outStrCalcK2)
    print(outStrCalcK2)
    document.add_paragraph(outStrCalcK4)
    print(outStrCalcK4)
    document.add_paragraph(outStrCalcK8)
    print(outStrCalcK8)
    document.add_paragraph(outStrCalcK16)
    print(outStrCalcK16)
    document.add_paragraph()
    print("")
    document.add_paragraph(outStrCalcFailBitPos)
    print(outStrCalcFailBitPos)
    outStrResult = f"Значення {posOfFailBit_} вказує на те, що помилка сталася в {int(posOfFailBit_, 2)} біті ({posOfFailBit_}={int(posOfFailBit_, 2)})."
    document.add_paragraph(outStrResult)
    print(outStrResult)  
    document.add_page_break()


def task1_6(pn, ll, document):
    print("\n\n1.6")
    document.add_paragraph("1.6 Для послідовності 16-кових цифр (1ц1л)(2ц1л)(1ц2л)(2ц2л)(1ц3л)(2ц3л)...(2ц7л)( 1ц8л)(2ц8л), користуючись картами Карно, визначити всі можливі помилкові коди, які можуть виникати при переході від цифри до цифри.")
    m = Graph(KARNAUGH_MAP)

    print("Персональні коди:")
    dtp = ""
    for i in range(len(pn)):
        dtp += f"{ll[i]} → {pn[i][0]}{pn[i][1]}  "
        print(f"{ll[i]} - {pn[i]} ", end='')
    print()
    document.add_paragraph(dtp)

    conversionList = [str(i) for i in list(chain.from_iterable(pn))]
    print(f"Список переходів, трансформований з персональних кодів:")
    document.add_paragraph("Визначення всіх можливих помилкових кодів:")
    dtp = print_with_arrows(conversionList)
    document.add_paragraph(dtp)
    table = tabulate(KARNAUGH_MAP_TABLE, tablefmt="grid")
    firstTable = document.add_table(rows=17, cols=6)
    #firstTable.autofit = False
    tableNumber = 1
    for i in range(17):
        for j in range(6):
            if j == 0 or j == 3:
                firstTable.rows[i].cells[j].width = Cm(0.5)
            elif i % 2 == 0 and (j == 1 or j == 4):
                    a = firstTable.cell(i, j)
                    b = firstTable.cell(i, j+1)
                    a.merge(b)
            if i % 2 != 0 and (j == 1 or j == 4) and tableNumber != 16:
                firstTable.rows[i].cells[j-1].text = str(tableNumber) + ")"
                tableNumber += 1
                cell = firstTable.rows[i].cells[j]
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture("images/karno.png", width = Cm(3.7), height = Cm(3.7))

    row = 1
    k = 0

    for i in range(15):
        splitter()
        print(table)
        codes_walked_through = ""
        error_codes = ""
        if conversionList[i] != conversionList[i + 1]:
            b1 = str(bin(int(conversionList[i], 16)))
            b2 = str(bin(int(conversionList[i + 1], 16)))
            b1 = b1.lstrip("0b").zfill(4)
            b2 = b2.lstrip("0b").zfill(4)
            conversionListText = f"{conversionList[i]} → {conversionList[i + 1]}: {b1} → {b2}\n"
            print(f"{conversionList[i]} → {conversionList[i + 1]}: {b1} → {b2}")
            n = m.all_shortest_paths(conversionList[i], conversionList[i + 1])
            for element in n:
                codes_walked_through += print_with_arrows(element) + "\n"

            if len(n) != 1:
                p = "Помилкові коди:"
                l = []
                for j in n:
                    z = j
                    del z[0]
                    del z [-1]
                    for element in z:
                        l.append(element)
                    l = delete_repeating_elements_from_list(l)
                for element in l:
                    b = str(bin(int(element, 16)))
                    b = b.lstrip("0b").zfill(4)
                    p += f" {b},"
                p = p.rstrip(",")
                print(p)
                error_codes = p
            else:
                print("Помилкових кодів немає.")
                error_codes = "Помилкових кодів немає"
        else:
            b1 = str(bin(int(conversionList[i], 16)))
            b2 = str(bin(int(conversionList[i + 1], 16)))
            b1 = b1.lstrip("0b").zfill(4)
            b2 = b2.lstrip("0b").zfill(4)
            conversionListText = f"{conversionList[i]} → {conversionList[i + 1]}: {b1} → {b2}\n"
            print(f"{conversionList[i]} → {conversionList[i + 1]}: {b1} → {b2}\n")
            print("Помилкових кодів немає")
            error_codes = "Помилкових кодів немає"
        if k == 2:
            row += 2
            k = 0
        if int(i+1) % 2 != 0:
            firstTable.rows[row+1].cells[2].text = error_codes
            paragraph = firstTable.rows[row+1].cells[2].paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size= Pt(8)
            firstTable.rows[row].cells[2].text = conversionListText + codes_walked_through
            paragraph = firstTable.rows[row].cells[2].paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size= Pt(8)
        else:
            firstTable.rows[row+1].cells[5].text = error_codes
            paragraph = firstTable.rows[row+1].cells[5].paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size= Pt(8)
            firstTable.rows[row].cells[5].text = conversionListText + codes_walked_through
            paragraph = firstTable.rows[row].cells[5].paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size= Pt(8)
        k += 1
    document.add_page_break()


def task1(personalNumbers, listedLetters, document):
    task1_1(personalNumbers, listedLetters, document)
    task1_2(personalNumbers, listedLetters, document)
    task1_3(personalNumbers, document)
    task1_4(listedLetters, personalNumbers, document)
    task1_5(listedLetters, personalNumbers, document)
    task1_6(personalNumbers, listedLetters, document)
