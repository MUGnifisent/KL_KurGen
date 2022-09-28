from utils.task2_utils import *                 #ver 0.1.3
from docx.shared import Cm, Pt
from utils.docx_utils import *

def task2_1(pn):
    print("\n\n2.1")

    print("Числа:")
    firstNumber = bin(int(str(pn[3][0]), 10))[2:]
    secondNumber = bin(int(str(pn[6][1]), 10))[2:]
    

    for i in range(4):
        if len(str(firstNumber)) < 4:
            firstNumber = "0" + str(firstNumber)
        if len(str(secondNumber)) < 4:
            secondNumber = "0" + str(secondNumber)

    print(f"1ц4л - {pn[3][0]} -> {firstNumber} \n2ц7л - {pn[6][1]} -> {secondNumber}\n")

    bothNumbers = str(firstNumber) + str(secondNumber)
    TZ2 = [[0, 0, 0, 0], [0, 0, 1, 0], [0, 1, 0, 0], [0, 1, 1, 0], [1, 0, 0, 0], [1, 0, 1, 0], [1, 1, 0, 0], [1, 1, 1, 0]]

    for i in range(8):
        TZ2[i][3] = int(str(bothNumbers)[i])

    print("a b c | f")
    for i in range(8):
        print(TZ2[i][0], TZ2[i][1], TZ2[i][2],"|", TZ2[i][3])

    Rule1 = False
    Rule2 = False

    if TZ2[0][3] == 0:
        print("1) Функція на нульовому наборі змінних f(0,0,0) = 0. Отже, функція зберігає константу «0».")
        Rule1 = True
    else:
        print("1) Функція на нульовому наборі змінних f(0,0,0) = 1. Отже, функція не зберігає константу «0».")
    if TZ2[7][3] == 1:
        print("2) Функція на одиничному наборі змінних f(1,1,1) = 1. Отже, функція зберігає константу «1».")
        Rule2 = True
    else:
        print("2) Функція на одиничному наборі змінних f(1,1,1) = 0. Отже, функція не зберігає константу «1».")

    monoTableArray = [[[0 for x in range(4)] for y in range(4)] for z in range(6)]


    for i in range(6):
        monoTableArray[i][0] = [TZ2[0][0], TZ2[0][1], TZ2[0][2], TZ2[0][3]]
    for i in range(2):
        monoTableArray[i][1] = [TZ2[1][0], TZ2[1][1], TZ2[1][2], TZ2[1][3]]
    for i in range(2):
        monoTableArray[i+2][1] = [TZ2[2][0], TZ2[2][1], TZ2[2][2], TZ2[2][3]]
    for i in range(2):
        monoTableArray[i+4][1] = [TZ2[4][0], TZ2[4][1], TZ2[4][2], TZ2[4][3]]
    monoTableArray[0][2] = [TZ2[3][0], TZ2[3][1], TZ2[3][2], TZ2[3][3]]
    monoTableArray[1][2] = [TZ2[5][0], TZ2[5][1], TZ2[5][2], TZ2[5][3]]
    monoTableArray[2][2] = [TZ2[3][0], TZ2[3][1], TZ2[3][2], TZ2[3][3]]
    monoTableArray[3][2] = [TZ2[6][0], TZ2[6][1], TZ2[6][2], TZ2[6][3]]
    monoTableArray[4][2] = [TZ2[5][0], TZ2[5][1], TZ2[5][2], TZ2[5][3]]
    monoTableArray[5][2] = [TZ2[6][0], TZ2[6][1], TZ2[6][2], TZ2[6][3]]
    for i in range(6):
        monoTableArray[i][3] = [TZ2[7][0], TZ2[7][1], TZ2[7][2], TZ2[7][3]]

    isMonotonic = True;

    for i in range(6):
        tempValue = -1
        for j in range(4):
            if monoTableArray[i][j][3] == 0 and tempValue == 1:
                tempValue = monoTableArray[i][j][3]
                monoTableArray[i][j][3] = 9 #Код нуля для правильного формування таблиці при її виводі
                monoTableArray[i][j-1][3] = 10 #Код одиниці для правильного формування таблиці при її виводі
                isMonotonic = False
            else:
                tempValue = monoTableArray[i][j][3]

    if isMonotonic:
        print("3) Функція є монотонною, оскільки при будь-якому зростанні кількості '1' у послідовності сусідніх наборів змінних значення функції не зменшується.")
    else:
        print("3) Функція не є монотонною, оскільки при будь-якому зростанні кількості '1' у послідовності сусідніх наборів змінних значення функції зменшується.")

    print("a b c | f    a b c | f    a b c | f    a b c | f    a b c | f    a b c | f")

    for i in range(4):
        for j in range(6):
            for k in range(4):
                if k != 3:
                    print(monoTableArray[j][i][k], end=" ")
                else:
                    if monoTableArray[j][i][k] == 9:
                        print("||o̲|", end="")
                    elif monoTableArray[j][i][k] == 10:
                        print("||̅₁|", end="")
                    else:
                        print("|", monoTableArray[j][i][k], end=" ")
            print(end="   ")
        print()
    
    isSelfdual = [0,0,0,0]
    Rule4 = False
    for i in range(4):
        if TZ2[i][3] == TZ2[7-i][3]:
            isSelfdual[i] = 1

    if isSelfdual[0] + isSelfdual[1] + isSelfdual[2] + isSelfdual[3] == 0:
        print("4) Функція є самодвоїстою")
        Rule4 = True
    else:
        print("4) Функція не є самодвоїстою, оскільки на", end=" ")
        if isSelfdual[0] == 1:
            print("першій,", end=" ")
        if isSelfdual[1] == 1:
          print("другій,", end=" ")
        if isSelfdual[2] == 1:
          print("третій,", end=" ")
        if isSelfdual[3] == 1:
            print("четвертій,", end=" ")       
        print("\b\b парі протилежних наборів функція не приймає протилежні значення.")
    print("a b c | f | a b c | f")
    for i in range(4):
        print(TZ2[i][0], TZ2[i][1], TZ2[i][2], "|", TZ2[i][3], "|", TZ2[7-i][0], TZ2[7-i][1], TZ2[7-i][2], "|", TZ2[7-i][3])


    print("5) Для визначення лінійності функції подамо її у вигляді полінома Жегалкіна:")
    Zhegalkin = "" 

    for i in range(8):
        if TZ2[i][3] == 1:
            if TZ2[i][0] == 0:
                Zhegalkin += "/a"
            else:
                Zhegalkin += "a"
            if TZ2[i][1] == 0:
                Zhegalkin += "/b"
            else:
                Zhegalkin += "b"
            if TZ2[i][2] == 0:
                Zhegalkin += "/c"
            else:
                Zhegalkin += "c"
            Zhegalkin += "⊕ "

    Zhegalkin = Zhegalkin[:-1]
    print(Zhegalkin, end="\b = ")
    
    ZhegalkinStep2 = ""
    ZhegalkinStep2Skip = False

    for i in range(len(Zhegalkin)):
        if ZhegalkinStep2Skip == False:
            if Zhegalkin[i] == "/":
                ZhegalkinStep2 += "(" + Zhegalkin[i+1] + " + 1)"
                ZhegalkinStep2Skip = True
            else:
                ZhegalkinStep2 += Zhegalkin[i]
        else:
            ZhegalkinStep2Skip = False
    

    for i in range(len(ZhegalkinStep2)):
        if ZhegalkinStep2[i] == "+":
            print("⊕ ", end="")
        elif i == len(ZhegalkinStep2) - 1:
            print(" = ", end="")
        else:
            print(ZhegalkinStep2[i], end="")
    
    

    ZhegalkinStep3Print = ""

    ZhegalkinStep3I = 0

    ZhegalkinFinalStep = []

    for step3 in range(ZhegalkinStep2.count('⊕')):
        isNumbersToWrite = False
        ZhegalkinStep3 = [[],[],[]]
        index = 0
        
        for i in range(ZhegalkinStep3I, len(ZhegalkinStep2)):
            ZhegalkinStep3I = i+1
            if ZhegalkinStep2[i] == "(" and isNumbersToWrite == False:
                isNumbersToWrite = True
            elif isNumbersToWrite == True and ZhegalkinStep2[i] != " " and ZhegalkinStep2[i] != "+" and ZhegalkinStep2[i] != ")":
                if ZhegalkinStep2[i-1] == "a" or ZhegalkinStep2[i-1] == "b" or ZhegalkinStep2[i-1] == "c" or ZhegalkinStep2[i-1] == "1":
                    ZhegalkinStep3[index][len(ZhegalkinStep3[index])-1] += ZhegalkinStep2[i]
                else:
                    ZhegalkinStep3[index].append(ZhegalkinStep2[i])
            elif ZhegalkinStep2[i] != "(" and ZhegalkinStep2[i] != ")" and ZhegalkinStep2[i] != "⊕" and ZhegalkinStep2[i] != " " and isNumbersToWrite == False:
                ZhegalkinStep3[index].append(ZhegalkinStep2[i])
                index+=1
            elif ZhegalkinStep2[i] == ")":
                isNumbersToWrite = False
                index+=1            
            elif index == 3:
                break
            elif ZhegalkinStep2[i] == "⊕":
                break

        if index == 3:
            ZhegalkinStep3Print += "("
            for i in range(len(ZhegalkinStep3[0])):
                for j in range(len(ZhegalkinStep3[1])):
                    for k in range(len(ZhegalkinStep3[2])):
                        if str(ZhegalkinStep3[0][i]) + str(ZhegalkinStep3[1][j]) + str(ZhegalkinStep3[2][k]) != "111":
                            ZhegalkinStep3Print += ''.join(sorted(str(ZhegalkinStep3[0][i]) + str(ZhegalkinStep3[1][j]) + str(ZhegalkinStep3[2][k]))).replace("1","") + " ⊕  "
                            ZhegalkinFinalStep.append(''.join(sorted(str(ZhegalkinStep3[0][i]) + str(ZhegalkinStep3[1][j]) + str(ZhegalkinStep3[2][k]))).replace("1",""))
                        else:
                            ZhegalkinStep3Print += "1 ⊕  "
                            ZhegalkinFinalStep.append("1")
            ZhegalkinStep3Print = ZhegalkinStep3Print[:-4]
            ZhegalkinStep3Print += ")⊕ "

    ZhegalkinStep3Print = ZhegalkinStep3Print[:-2]
    print(ZhegalkinStep3Print, "= ")
    print("= ", end="")
    ZhegalkinFinalOutput = []
    ZhegalkinFilter = []
    SukaYakYaZaebavsya = []
    for i in range(len(ZhegalkinFinalStep)):
        ZhegalkinOutputLine = "" #вибачте за такі круті назви змінних, моя голова вмерла
        count = 0
        ZhegalkinFilter.append(ZhegalkinFinalStep[i])
        ZhegalkinFinalOutput.append(ZhegalkinFinalStep[i])
        for k in range(len(ZhegalkinFilter)):
            if ZhegalkinFinalStep[i] == ZhegalkinFilter[k]:
                count += 1
        if count == 1:
            count = 0
            if i != 0:
                ZhegalkinOutputLine += "⊕ "
            for j in range(len(ZhegalkinFinalStep)):
                if ZhegalkinFinalStep[i] == ZhegalkinFinalStep[j]:
                    ZhegalkinOutputLine += str(ZhegalkinFinalStep[j]) + "⊕  "
                    count += 1
            if count % 2 == 0:
                ZhegalkinOutputLine += "    (= 0 )"
                ZhegalkinFinalOutput.pop(len(ZhegalkinFinalOutput)-1)
            else:
                ZhegalkinOutputLine += "    (=" + ZhegalkinFinalStep[i] + ")"
        else:
            ZhegalkinFilter.pop(len(ZhegalkinFilter)-1)
            ZhegalkinFinalOutput.pop(len(ZhegalkinFinalOutput)-1)
        if len(ZhegalkinOutputLine) != 0:
            SukaYakYaZaebavsya.append(ZhegalkinOutputLine)

    for i in range(len(SukaYakYaZaebavsya)):
        if i == len(SukaYakYaZaebavsya) - 1:
            for j in range(len(SukaYakYaZaebavsya[i])):
                if SukaYakYaZaebavsya[i][len(SukaYakYaZaebavsya[i]) - 1 - j] == "⊕":
                    list1 = list(SukaYakYaZaebavsya[i])
                    list1[len(SukaYakYaZaebavsya[i]) - 1 - j] = "="
                    SukaYakYaZaebavsya[i] = '' .join(list1)
                    break
        print(SukaYakYaZaebavsya[i])
    
    print("= ",end="")
    remains = "("
    for i in range(len(ZhegalkinFinalOutput)):
        if i == len(ZhegalkinFinalOutput)-1:
            print(ZhegalkinFinalOutput[i])
        else:
            print(ZhegalkinFinalOutput[i], "⊕  " ,end="")
        if len(ZhegalkinFinalOutput[i]) != 1:
            remains += str(ZhegalkinFinalOutput[i]) + ", "


    

    Rule5 = False

    if len(remains) > 2:
        print(remains, "\b\b - це добутки змінних)")
        print("Оскільки поліном містить добутки змінних, то функція нелінійна.")
    else:
        print("Оскільки поліном не містить добутки змінних, то функція лінійна.")
        Rule5 = True

    rulesCount = 0

    rules = ""

    if Rule1:
        rulesCount += 1
        rules += "не зберігання константи '0', "
    if Rule2:
        rulesCount += 1
        rules += "не зберігання константи '1', "
    if isMonotonic:
        rulesCount += 1
        rules += "немонотонність, "
    if Rule4:
        rulesCount += 1
        rules += "несамодвоїсність, "
    if Rule5:
        rulesCount += 1
        rules += "нелінійність, "

    print()
    #if rulesCount == 1:
    #    print("Отже, із п'яти необхідних для створення ФПС властивостей відсутня одна ", end="")
    #elif rulesCount == 2:
    #    print("Отже, із п'яти необхідних для створення ФПС властивостей відсутні дві ", end="")
    #elif rulesCount == 3:
    #    print("Отже, із п'яти необхідних для створення ФПС властивостей відсутні три ", end="")
    #elif rulesCount == 4:
    #    print("Отже, із п'яти необхідних для створення ФПС властивостей відсутні чотири ", end="")
    #elif rulesCount == 5:
    #    print("Отже, із п'яти необхідних для створення ФПС властивостей відсутні п'ять ", end="")
    #elif rulesCount == 0:
    #    print("Отже, із п'яти необхідних для створення ФПС властивостей присутні усі", end="")

    match rulesCount:
        case 1:
            print("Отже, із п'яти необхідних для створення ФПС властивостей відсутня одна ", end="")
        case 2:
            print("Отже, із п'яти необхідних для створення ФПС властивостей відсутні дві ", end="")
        case 3:
            print("Отже, із п'яти необхідних для створення ФПС властивостей відсутні три ", end="")
        case 4:
            print("Отже, із п'яти необхідних для створення ФПС властивостей відсутні чотири ", end="")
        case 5:
            print("Отже, із п'яти необхідних для створення ФПС властивостей відсутні п'ять ", end="")
        case 0:
            print("Отже, із п'яти необхідних для створення ФПС властивостей присутні усі", end="")

    if rulesCount != 0:
        print(rules, "\b\b тому дана функція не утворює ФПС.")
    else:
        print(" тому дана функція утворює ФПС.")


def task2_2(personalNumbers, listedLetters, document):
    print('\n\n2.2')
    document.add_paragraph("""2.2 Мінімізувати за допомогою методу Квасна-Мак-Класкі-Петрика 5 функцій (f0, f1, f2, f3, f4) 5-ти змінних (a, b, c, d, e). Функції задано за допомогою таблиці ТZ.3. Побудувати таблицю, яка ілюструє процес знаходження простих імплікант,і таблицю покриття (імплікантну таблицю). За допомогою методу Петрика визначити всі мінімальні розв'язки. Кожний третій набір для кожної з функцій має невизначене значення. Відлік починається від першого згори одиничного значення функції з врахуванням зміщення, величина якого позначена у табл. ТZ.3:
    S - відлік починається безпосередньо від першого згори одиничного значення;
    -S - відлік починається від попереднього набору відносно першого згори одиничного значення;
    +S - відлік починається від наступного набору відносно першого згори одиничного значення.""")
    pn = personalNumbers

    LstOfBinNum = []
    for i in range(8):
        Num1Bin = bin(pn[i][0])[2:]
        LstOfBinNum.append(Num1Bin.zfill(4))
        Num2Bin = bin(pn[i][1])[2:]
        LstOfBinNum.append(Num2Bin.zfill(4))

    firstTable = document.add_table(rows=6, cols=16, style="Table Grid")
    firstTable.autofit = False

    mergeindex = 0

    for i in range(8):
        a = firstTable.cell(0, mergeindex)
        mergeindex+=1
        b = firstTable.cell(0, mergeindex)
        mergeindex+=1
        a.merge(b)
        cell = firstTable.rows[0].cells[mergeindex-1]
        cell.text = str(listedLetters[i])
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size= Pt(11)

    print(' ', end=' ')
    print(*listedLetters, sep=' '*7)
    pa = CreateTableOfBinNum(LstOfBinNum, pn)
    print(pa)
    for i in range(len(pa)):
        for j in range(len(pa[0])):
            cell = firstTable.rows[i+1].cells[j]
            cell.text = str(pa[i][j])
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(11)

    print('\n\n')
    document.add_paragraph("")
    secondTable = document.add_table(rows=34, cols=11, style="Table Grid")
    secondTable.autofit = False
    for i in range(34):
        for j in range(11):
            cell = secondTable.rows[i].cells[j]
            if j == 0:
                cell.width = Cm(0.8)
            elif j > 0 and j < 6:
                cell.width = Cm(0.7)
            elif j > 5:
                cell.width = Cm(1.1)

    a = secondTable.cell(0, 0)
    b = secondTable.cell(0, 5)
    a.merge(b)
    secondTable.rows[0].cells[0].text = 'Зміщення першого невизначеного значення'
    secondTable.rows[0].cells[6].text = '-S'
    secondTable.rows[0].cells[7].text = 'S'
    secondTable.rows[0].cells[8].text = '+S'
    secondTable.rows[0].cells[9].text = '-S'
    secondTable.rows[0].cells[10].text = 'S'
    secondTable.rows[1].cells[0].text = '№'
    secondTable.rows[1].cells[1].text = 'e'
    secondTable.rows[1].cells[2].text = 'd'
    secondTable.rows[1].cells[3].text = 'c'
    secondTable.rows[1].cells[4].text = 'b'
    secondTable.rows[1].cells[5].text = 'a'
    secondTable.rows[1].cells[6].text = 'f₀'
    secondTable.rows[1].cells[7].text = 'f₁'
    secondTable.rows[1].cells[8].text = 'f₂'
    secondTable.rows[1].cells[9].text = 'f₃'
    secondTable.rows[1].cells[10].text = 'f₄'
    for i in range(5):
        secondTable.rows[0].cells[i+6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        secondTable.rows[0].cells[i+6].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i in range(11):
        secondTable.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    secondTableData = [[] for x in range(32)]
    print('Зміщення першого\nневизначеного\t   \t-S\tS\t+S\t-S\tS\nзначення')
    print('№ набору e d c b a\tf0\tf1\tf2\tf3\tf4')
    BinNum0to31 = []
    NumF0 = []
    letter = -1
    SetsForF0 = []
    StartPosXInF0, StartPosXInF1, StartPosXInF2, StartPosXInF3, StartPosXInF4 = FindPosXInF(LstOfBinNum)
    for i in range(32):
        print(i,end='\t ')
        secondTableData[i].append(i)
        BinNum0to31.append(bin(i)[2:].zfill(5))
        for number in BinNum0to31[i]:
            secondTableData[i].append(number)
            print(number, end=' ')
        
        if i%4 == 0:
            letter += 1
        print(end='\t')
        if i == StartPosXInF0:
             NumF0.append('x')
             secondTableData[i].append(f'({LstOfBinNum[letter][i%4]})x')
             print(f'({LstOfBinNum[letter][i%4]})x',end='\t')
             StartPosXInF0 += 3
        else:
            NumF0.append(LstOfBinNum[letter][i%4])
            secondTableData[i].append(LstOfBinNum[letter][i%4])
            print(LstOfBinNum[letter][i%4], end='\t')
        d = letter + 8
        StartPosXInF1, temp = FillFOfX(StartPosXInF1, LstOfBinNum, letter, i)
        secondTableData[i].append(temp)
        StartPosXInF2, temp = FillFOfX(StartPosXInF2, LstOfBinNum, letter, i)
        secondTableData[i].append(temp)
        StartPosXInF3, temp = FillFOfX(StartPosXInF3, LstOfBinNum, d, i)
        secondTableData[i].append(temp)
        StartPosXInF4, temp = FillFOfX(StartPosXInF4, LstOfBinNum, d, i)
        secondTableData[i].append(temp)
        if NumF0[i] == '1' or NumF0[i] == 'x':
            SetsForF0.append(BinNum0to31[i])
        print()

    for i in range(len(secondTableData)):
        for j in range(len(secondTableData[i])):
            cell = secondTable.rows[i+2].cells[j]
            cell.text = str(secondTableData[i][j])
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(10)

    document.add_paragraph("")
    print('\n\n')

    abc = 97
    SortDict = SortInDictBy1(SetsForF0, abc)
    ListMerge, list1, listname = CreateListOfMerges(SortDict)
    AllDataForTable = []
    AllDataForTable.append(list1)
    AllDataForTable.append(listname)
    print('К\tП\tУ', end='\t')
    while True:
        if list(filter(lambda x: x[2] == '+', list1)) == []:
            break
        print('С\tК\tП\tУ', end='\t')
        abc += len(SortDict)
        SortDict = SortInDictBy1(ListMerge, abc)
        ListMerge, list1, listname = CreateListOfMerges(SortDict)
        AllDataForTable.append(list1)
        AllDataForTable.append(listname)
    print()
    MaxLen = len(max(AllDataForTable, key = lambda x: len(x)))

    thirdTableCols = 0

    print(AllDataForTable)
    for index in range(MaxLen):
        for table in AllDataForTable:
            if index == 0 and index < len(table):
                thirdTableCols += len(table[index])
            if index < len(table):
                print(*table[index], sep='\t', end='\t')
            else:
                print('\t\t', end='\t')
        print()

    document.add_paragraph("")
    thirdTable = document.add_table(rows=MaxLen+1, cols=thirdTableCols)
    thirdTable.rows[0].cells[0].text = 'К'
    thirdTable.rows[0].cells[1].text = 'П'
    thirdTable.rows[0].cells[2].text = 'У'
    thirdTableUpperLetters = 3
    for i in range(int((thirdTableCols-3)/4)):
        thirdTable.rows[0].cells[thirdTableUpperLetters].text = 'С'
        thirdTableUpperLetters += 1
        thirdTable.rows[0].cells[thirdTableUpperLetters].text = 'К'
        thirdTableUpperLetters += 1
        thirdTable.rows[0].cells[thirdTableUpperLetters].text = 'П'
        thirdTableUpperLetters += 1
        thirdTable.rows[0].cells[thirdTableUpperLetters].text = 'У'
        thirdTableUpperLetters += 1

    for i in range(thirdTableCols):
        block = thirdTable.rows[0].cells[i]
        paragraph = block.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs
        font = run[0].font
        font.size= Pt(10)
        set_cell_border(
            block,
            top={"sz": 5, "val": "single", "color": "#000000"},
            bottom={"sz": 5, "val": "single", "color": "#000000"},
            start={"sz": 5, "val": "single", "color": "#000000"},
            end={"sz": 5, "val": "single", "color": "#000000"},
        )

    row = 1
    col = 0
    for index in range(MaxLen):
        for table in AllDataForTable:
            if index < len(table):
                for cell in table[index]:
                    if col != thirdTableCols:
                        block = thirdTable.rows[row].cells[col]
                        block.text = cell
                        paragraph = block.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.runs
                        font = run[0].font
                        font.size= Pt(10)
                        col += 1
                        set_cell_border(
                            block,
                            top={"sz": 5, "val": "single", "color": "#000000"},
                            bottom={"sz": 5, "val": "single", "color": "#000000"},
                            start={"sz": 5, "val": "single", "color": "#000000"},
                            end={"sz": 5, "val": "single", "color": "#000000"},
                        )
            else:
                col += 3
        col = 0
        row += 1
    
    thirdTable.autofit = False
    for i in range(MaxLen+1):
        thirdTable.rows[i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        thirdTable.rows[i].height = Pt(13)
        for j in range(thirdTableCols):
            cell = thirdTable.rows[i].cells[j]
            if thirdTable.rows[0].cells[j].text == "К":
                cell.width = Cm(1.5)
            elif thirdTable.rows[0].cells[j].text == "П":
                cell.width = Cm(1)
            elif thirdTable.rows[0].cells[j].text == "У":
                cell.width = Cm(0.7)
            elif thirdTable.rows[0].cells[j].text == "С":
                cell.width = Cm(1.2)
    
    print('\n\n')
    AllNegative = []
    for i in range(0,len(AllDataForTable), 2):
        TempNegative = list(filter(lambda x: x[2] == '-', AllDataForTable[i]))
        for item in TempNegative:
            if item[0] not in AllNegative:
                AllNegative.append(item[0])
    document.add_paragraph("")
    document.add_paragraph("Прості імпліканти:")
    document.add_paragraph(', '.join(AllNegative))
    print('Прості імпліканти:')
    print(*AllNegative, sep = ', ', end='.\n\n')
    ListCountLetter = [len(x.replace('-','')) for x in AllNegative]
    LstF0Only1 = [BinNum0to31[i] for i in range(32) if NumF0[i] == '1']
    ListsWithAplicant = []
    ListOfLetters = {f'I{i+1}':BinaryToLetters(AllNegative[i]) for i in range(len(AllNegative))}
    print('Одиничні', end='\t')
    print(*[item.ljust(5) for item in ListOfLetters.keys()], sep='\t')
    print('набори', end='\t\t')
    print(*[item.ljust(5) for item in ListOfLetters.values()], sep='\t')
    print('edcba', end='\t\t')
    print(*AllNegative, sep='\t')
    fourthTable = document.add_table(rows=len(LstF0Only1)+4, cols=len(ListOfLetters.values())+1, style="Table Grid")
    a = fourthTable.cell(0, 0)
    b = fourthTable.cell(1, 0)
    a.merge(b)
    fourthTable.rows[0].cells[0].text = "Одиничні набори"
    fourthTable.rows[2].cells[0].text = "edcba"
    print(str(ListCountLetter))
    for i in range(len(ListOfLetters.keys())):
        fourthTable.rows[0].cells[i+1].text = "I" + conver_to_lower_symbol(str(i+1))
        p = fourthTable.rows[1].cells[i+1].add_paragraph()
        p._element.append(BinaryToLettersWithSpecialSymbols(AllNegative[i]))
        fourthTable.rows[2].cells[i+1].text = AllNegative[i]
        fourthTable.rows[len(LstF0Only1)+3].cells[i+1].text = str(ListCountLetter[i])
    row = 3
    col = 0
    for Set in LstF0Only1:
        col = 0
        fourthTable.rows[row].cells[col].text = Set
        col += 1
        print(Set, end='\t\t')
        TempList = []
        for merge in range(len(AllNegative)):
            if CheckEqualityNums(Set,AllNegative[merge]):
                print('V', end = '\t')
                TempList.append(f'I{merge + 1}')
                fourthTable.rows[row].cells[col].text = "✓"
            else:
                fourthTable.rows[row].cells[col].text = " "
                print('-', end = '\t')
            col += 1
        ListsWithAplicant.append(TempList)
        print('\n')
        row += 1
    fourthTable.rows[len(LstF0Only1)+3].cells[0].text = "Літер в імплакації"
    for i in range(len(LstF0Only1)+4):
        for j in range(len(ListOfLetters.values())+1):
            block = fourthTable.rows[i].cells[j]
            if i == 1 and j > 0:
                print(f"i {i} j {j}")
            else:
                paragraph = block.paragraphs[0]
                run = paragraph.runs
                font = run[0].font
                font.size= Pt(10)


    print('Літер в\nімплакації', end='\t')
    print(*ListCountLetter, sep='\t')
    document.add_paragraph("")
    print('\n\n')
    DictOfCountLetters = {f'I{index + 1}':ListCountLetter[index] for index in range(len(ListCountLetter))}
    document.add_paragraph("Отримаємо функцію F:")
    print('Отримаємо функцію F:\nF =',end=' ')
    document.add_paragraph("F = " + PrintFunctionAmplicants(ListsWithAplicant, ' v ', ' '))
    document.add_paragraph("Спрощення:")
    print('\nСпрощення:\nF =', end = ' ')
    
    ListOneAmplicant, ListAmplicants, ListResults = SimplifyingExpression(DictOfCountLetters, ListsWithAplicant)
    ListOneAmplicant = [''.join(i)  for i in ListOneAmplicant]
    for i in ListResults:
        i.update(ListOneAmplicant)
    ListOneAmplicant = sorted(ListOneAmplicant, key = lambda x: int(x[1:]))
    ListAmplicants = [sorted(list(i), key = lambda x: int(x[1:])) for i in ListAmplicants]
    ListResults = [sorted(list(i), key = lambda x: int(x[1:])) for i in ListResults]
    ListOneAmplicantOutput = ""
    for i in range(len(ListOneAmplicant)):
        for j in range(len(ListOneAmplicant[i])):
            if ListOneAmplicant[i][j].isdigit():
                ListOneAmplicantOutput += conver_to_lower_symbol(ListOneAmplicant[i][j])
            else:
                ListOneAmplicantOutput += ListOneAmplicant[i][j]
        ListOneAmplicantOutput += " "
    
    print(*ListOneAmplicant, sep=' ', end='')
    ListOneAmplicantOutput += PrintFunctionAmplicants(ListAmplicants, ' v ', ' ')
    ListOneAmplicantOutput += " = "
    print(' = ', end='')
    ListOneAmplicantOutput += PrintFunctionAmplicants(ListResults, ' ', ' v ')
    document.add_paragraph("F = " + ListOneAmplicantOutput)
    document.add_paragraph(f'\nОтже, дана функція F має {len(ListResults)} мінімальні ДНФ:')
    print(f'\n\nОтже, дана функція F має {len(ListResults)} мінімальні ДНФ:')
    for i, dnf in enumerate(ListResults, 1):
        temp = f'{i}) F = '
        firstPart = str(' v '.join(dnf))
        for z in range(len(firstPart)):
            for j in range(len(firstPart[z])):
                if firstPart[z][j].isdigit():
                    temp += conver_to_lower_symbol(firstPart[z][j])
                else:
                    temp += firstPart[z][j]
        p = document.add_paragraph()
        p.text = temp
        temp = ""
        temp += ' = ' + str(' v '.join([ListOfLetters[item] for item in dnf])) + ";"
        List0 = ['ē', 'đ', 'č', 'ƀ', 'ā']
        List1 = ['e', 'd', 'c', 'b', 'a']
        for z in range(len(temp)):
            index = 0
            for j in range(len(List0)):
                if temp[z] == List0[j]:
                    p._element.append(latex_to_word("\\bar{" + List1[j] + "}"))
                    break
                index +=1
                if index == 5:
                    if temp[z] == " ":
                        p._element.append(latex_to_word("\\medspace"))
                    elif temp[z] == "v":
                        p._element.append(latex_to_word("\\text{v}"))
                    else:
                        p._element.append(latex_to_word(temp[z]))
        
        print(f'{i}) F = ', end ='')
        print(*[item.ljust(3) for item in dnf], sep=' v ', end =' = ')
        print(*[ListOfLetters[item] for item in dnf], sep=' v ', end =';\n')


def task2(personalNumbers, listedLetters, document):
    task2_1(personalNumbers)
    task2_2(personalNumbers, listedLetters, document)
