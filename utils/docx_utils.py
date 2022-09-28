from datetime import date
from docx.shared import Pt

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_TABLE_ALIGNMENT


def title_page(document, PIB):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    document.add_paragraph('Міністерство освіти і науки України').alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('Національний університет "Львівська Політехніка"').alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('Кафедра ЕОМ').alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('\n\n\n\n\n\n\nЗвіт').alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('до курсової роботи').alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('з дисципліни "Комп\'ютерна логіка"').alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('\n\n\n\n\n\nВиконав:').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph('ст. групи КІ-210').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph(PIB.split(' ', 1)[0] + " " + PIB[PIB.find(" ",0)+1] + ". " + PIB[PIB.find(" ", PIB.find(" ",0)+1)+1] + ".").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph('Прийняв:').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph('доцент каф. ЕОМ').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph('Голембо В. А.').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph('\n\n\n\n\n\n\n\n\nЛьвів ' + str(date.today().year)).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))